import shutil
import sys
import os
from datetime import datetime  # 添加这行导入语句
import random  # 添加随机数模块
import pandas as pd


# IMPORT / GUI AND MODULES AND WIDGETS
# ///////////////////////////////////////////////////////////////
from modules import *
from widgets import *
from openpyxl import load_workbook
from PySide6.QtCharts import (QChart, QLineSeries, QValueAxis, QChartView, 
                             QBarSeries, QBarSet, QBarCategoryAxis, QPieSeries,
                             QPieSlice)
import sqlite3
import threading
from PySide6.QtWidgets import (QComboBox, QLabel, QVBoxLayout, QHBoxLayout, 
                             QWidget, QCheckBox, QListView, QPushButton, 
                             QFrame, QScrollArea, QTextEdit, QFileDialog,
                             QMainWindow, QGridLayout, QCalendarWidget)
from PySide6.QtCore import Qt, QSize, QDate
from PySide6.QtGui import (QPixmap, QCursor, QWheelEvent, QPainter, QColor, 
                          QPen, QBrush, QFont, QStandardItemModel, QStandardItem)
from modules.OtherView.CustomMessagebox import CustomMessageBox
from modules.Script.genReportRTemplate import generate_report_with_images



os.environ["QT_FONT_DPI"] = "96"  # FIX Problem for High DPI and Scale above 100%

# SET AS GLOBAL WIDGETS
# ///////////////////////////////////////////////////////////////
widgets = None


class ImageViewer(QMainWindow):
    def __init__(self, image_path):
        super().__init__()
        self.scale_factor = 1.0

        # 创建滚动区域
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setAlignment(Qt.AlignmentFlag.AlignCenter)  # 让图片居中
        self.setCentralWidget(self.scroll_area)

        # 创建标签来显示图片
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.scroll_area.setWidget(self.image_label)

        # 加载图片
        self.pixmap = QPixmap(image_path)

        # 计算窗口和图片的合适大小
        screen_size = QApplication.primaryScreen().size()
        max_width = int(screen_size.width() * 0.8)
        max_height = int(screen_size.height() * 0.8)

        img_width = self.pixmap.width()
        img_height = self.pixmap.height()

        # 计算缩放比例
        scale_ratio = min(max_width / img_width, max_height / img_height, 1.0)
        display_width = int(img_width * scale_ratio)
        display_height = int(img_height * scale_ratio)
        self.scale_factor = scale_ratio

        # 缩放图片（如果需要）
        if scale_ratio < 1.0:
            scaled_pixmap = self.pixmap.scaled(display_width, display_height, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            self.image_label.setPixmap(scaled_pixmap)
        else:
            self.image_label.setPixmap(self.pixmap)

        # 设置窗口大小
        self.resize(display_width, display_height)

    def wheelEvent(self, event: QWheelEvent):
        # 处理鼠标滚轮事件
        if event.angleDelta().y() > 0:
            self.scale_factor *= 1.1  # 放大10%
        else:
            self.scale_factor *= 0.9  # 缩小10%

        # 限制缩放范围
        self.scale_factor = max(0.1, min(self.scale_factor, 10.0))

        # 计算新的图片大小
        new_width = int(self.pixmap.width() * self.scale_factor)
        new_height = int(self.pixmap.height() * self.scale_factor)

        # 更新图片显示
        scaled_pixmap = self.pixmap.scaled(new_width, new_height,
                                           Qt.AspectRatioMode.KeepAspectRatio,
                                           Qt.TransformationMode.SmoothTransformation)
        self.image_label.setPixmap(scaled_pixmap)
        self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.scroll_area.setAlignment(Qt.AlignmentFlag.AlignCenter)


class ImageGallery(QFrame):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.image_viewer = None
        self.init_ui()
        
    def init_ui(self):
        # 创建网格布局
        self.grid_layout = QGridLayout(self)
        self.grid_layout.setSpacing(10)
        self.grid_layout.setContentsMargins(10, 10, 10, 10)
        
    def load_images(self, sku: str):
        """加载指定SKU的图纸
        
        Args:
            sku: SKU编号
        """
        # 清除现有图片
        for i in reversed(range(self.grid_layout.count())): 
            self.grid_layout.itemAt(i).widget().setParent(None)
            
        # 获取图纸信息
        db_manager = DatabaseManager()
        image_infos = db_manager.get_drawing_info_by_sku(sku)
        
        # 计算每行显示的图片数量
        max_cols = 4  # 每行最多显示4张图片
        total_images = len(image_infos)
        if total_images > 0:
            # 计算每行实际显示的图片数量
            actual_cols = min(max_cols, total_images)
            # 计算缩略图大小
            thumbnail_size = min(150, 800 // actual_cols - 20)  # 800是容器宽度，20是间距
            
            # 显示图片
            row = 0
            col = 0
            
            for image_info in image_infos:
                image_path = image_info["drawing_path"]
                if os.path.exists(image_path):
                    try:
                        # 创建容器widget
                        container = QWidget()
                        container_layout = QVBoxLayout(container)
                        container_layout.setContentsMargins(0, 0, 0, 0)
                        container_layout.setSpacing(5)
                        
                        # 创建缩略图标签
                        thumbnail_label = QLabel()
                        pixmap = QPixmap(image_path)
                        if not pixmap.isNull():
                            # 设置缩略图大小
                            thumbnail_label.setPixmap(pixmap.scaled(thumbnail_size, thumbnail_size, 
                                                                 Qt.AspectRatioMode.KeepAspectRatio,
                                                                 Qt.TransformationMode.SmoothTransformation))
                            thumbnail_label.setCursor(Qt.CursorShape.PointingHandCursor)
                            # 添加点击事件
                            thumbnail_label.mousePressEvent = lambda e, path=image_path: self.show_full_image(path)
                            
                            # 创建图片名称标签
                            name_label = QLabel()
                            name_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                            name_label.setStyleSheet("font-size: 10px;")
                            name_label.setText(f"{image_info['word_part']} {image_info['part_number']}")
                            
                            # 添加到容器布局
                            container_layout.addWidget(thumbnail_label)
                            container_layout.addWidget(name_label)
                            
                            # 添加到网格布局
                            self.grid_layout.addWidget(container, row, col)
                            
                            # 更新行列位置
                            col += 1
                            if col >= actual_cols:
                                col = 0
                                row += 1
                    except Exception as e:
                        CustomMessageBox.warning(None, "警告", f"加载图片时出错: {str(e)}")

    def load_real_thing_images(self, sku: str):
        print("load_real_thing_images")
        """加载指定SKU的实物图片
        
        Args:
            sku: SKU编号
        """
        # 清除现有图片
        for i in reversed(range(self.grid_layout.count())): 
            self.grid_layout.itemAt(i).widget().setParent(None)
            
        # 获取实物图片信息
        db_manager = DatabaseManager()
        image_infos = db_manager.get_part_real_images_by_sku(sku)
        
        # 计算每行显示的图片数量
        max_cols = 4  # 每行最多显示4张图片
        total_images = len(image_infos)
        if total_images > 0:
            # 计算每行实际显示的图片数量
            actual_cols = min(max_cols, total_images)
            # 计算缩略图大小
            thumbnail_size = min(150, 800 // actual_cols - 20)  # 800是容器宽度，20是间距
            
            # 显示图片
            row = 0
            col = 0
            
            for image_info in image_infos:
                image_path = image_info["image_path"]
                if os.path.exists(image_path):
                    try:
                        # 创建容器widget
                        container = QWidget()
                        container_layout = QVBoxLayout(container)
                        container_layout.setContentsMargins(0, 0, 0, 0)
                        container_layout.setSpacing(5)
                        
                        # 创建缩略图标签
                        thumbnail_label = QLabel()
                        pixmap = QPixmap(image_path)
                        if not pixmap.isNull():
                            # 设置缩略图大小
                            thumbnail_label.setPixmap(pixmap.scaled(thumbnail_size, thumbnail_size, 
                                                                 Qt.AspectRatioMode.KeepAspectRatio,
                                                                 Qt.TransformationMode.SmoothTransformation))
                            thumbnail_label.setCursor(Qt.CursorShape.PointingHandCursor)
                            # 添加点击事件
                            thumbnail_label.mousePressEvent = lambda e, path=image_path: self.show_full_image(path)
                            
                            # 创建图片名称标签
                            name_label = QLabel()
                            name_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                            name_label.setStyleSheet("font-size: 10px;")
                            name_label.setText(f"{image_info['part_number']}")
                            
                            # 添加到容器布局
                            container_layout.addWidget(thumbnail_label)
                            container_layout.addWidget(name_label)
                            
                            # 添加到网格布局
                            self.grid_layout.addWidget(container, row, col)
                            
                            # 更新行列位置
                            col += 1
                            if col >= actual_cols:
                                col = 0
                                row += 1
                    except Exception as e:
                        CustomMessageBox.warning(None, "警告", f"加载图片时出错: {str(e)}")
                    
    def show_full_image(self, image_path):
        """显示完整图片"""
        self.image_viewer = ImageViewer(image_path)
        self.image_viewer.show()
    
    def clear_images(self):
        """清空所有图片"""
        for i in reversed(range(self.grid_layout.count())): 
            widget = self.grid_layout.itemAt(i).widget()
            if widget:
                widget.setParent(None)


class MainWindow(QMainWindow):
    
    def __init__(self):
        QMainWindow.__init__(self)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        # 全局变量
        self.selectedSku=""
        self.db_path = ProjectSettings.DATABASE_PATH
        # 创建数据库管理器实例
        self.db_manager = DatabaseManager()
        # 引入外部窗口
        self.ui_addSku = UI_AddSkusView()
        # 连接添加SKU窗口的信号到刷新列表的方法
        self.ui_addSku.sku_added.connect(self.reflesh_Skulist)
        # 连接删除信号
        self.ui.skuList.item_deleted.connect(self.delete_sku_from_database)
        
        self.camera_window = None  # 初始化相机窗口变量
        global widgets
        widgets = self.ui

        # 创建折线图
        self.create_line_chart()
        # 创建柱状图
        self.create_bar_chart()
        # 创建-饼图-机种分类
        self.create_pie_chart()
        # 创建日历视图
        # self.create_calendar()

        # USE CUSTOM TITLE BAR | USE AS "False" FOR MAC OR LINUX
        # ///////////////////////////////////////////////////////////////
        Settings.ENABLE_CUSTOM_TITLE_BAR = True

        # APP NAME
        # ///////////////////////////////////////////////////////////////
        title = "FPP E-FAI报告平台"
        description = "FPP E-FAI报告平台"
        # APPLY TEXTS
        self.setWindowTitle(title)
        widgets.titleRightInfo.setText(description)

   

        # SET UI DEFINITIONS
        # ///////////////////////////////////////////////////////////////
        UIFunctions.uiDefinitions(self)
        UIFunctions.toggleMenu(self, True)

        # BUTTONS CLICK
        # ///////////////////////////////////////////////////////////////

        # LEFT MENUS
        widgets.btn_home.clicked.connect(self.buttonClick)
        widgets.btn_new.clicked.connect(self.buttonClick)
        widgets.btn_save.clicked.connect(self.buttonClick)
        widgets.btn_Flow.clicked.connect(self.buttonClick)
        
        # 添加SKu，打开新的窗口，两种方法，单个上传，批量上传
        widgets.btn_AddMoreSku.clicked.connect(self.addMoreSku)
        # 搜索输入框，就更新列表
        widgets.skuSearchLine.textChanged.connect(self.filter_list)

        self.ui.skuList.itemClicked.connect(self.sku_clicked)
        
        
        self.ui.btnFlow.bom_check.clicked.connect(self.btnClick_BomCheck)
        self.ui.btnFlow.pic_download.clicked.connect(self.btnClick_PicDownload)
        self.ui.btnFlow.pic_check.clicked.connect(self.btnClick_PicCheck)
        self.ui.btnFlow.gen_report.clicked.connect(self.btnClick_GenReport)
        
        self.ui.bom_return.mousePressEvent = lambda event: self.btnClick_BomReturn()
        self.ui.pic_return.mousePressEvent = lambda event: self.btnClick_PicReturn()
        self.ui.genBack.mousePressEvent = lambda event: self.btnClick_genBack()
        
    
        self.ui.uploadCheckBtn.clicked.connect(self.uploadCheckFile)
        

        # 连接GenReport页面的信号
        self.ui.model_comboBox.currentTextChanged.connect(self.on_model_changed)
        self.ui.buildBtn.clicked.connect(self.generate_report)
        
        # 连接SNBtn的点击事件
        self.ui.SNBtn.clicked.connect(self.update_sn_label)

        self.ui.drawEndBtn.clicked.connect(self.drawEnd)
        self.ui.bomEndBtn.clicked.connect(self.bomEnd)
        # 窗口居中显示
        self.center_window()
        
        
        # sku查询
        self.ui.redoBtn.clicked.connect(self.redoSkuTable)
        self.ui.skuQueryBtn.clicked.connect(self.skuTableQuery)

        # 下载指定模板
        self.ui.downloadMBBtn.clicked.connect(self.downloadMB)

        # EXTRA LEFT BOX
        def openCloseLeftBox():
            UIFunctions.toggleLeftBox(self, True)

        # SHOW APP
        # ///////////////////////////////////////////////////////////////
        self.show()

        # SET CUSTOM THEME
        # ///////////////////////////////////////////////////////////////
        # 路径冻结，防止打包成exe后路径错乱
        if getattr(sys, 'frozen', False):
            absPath = os.path.dirname(os.path.abspath(sys.executable))
        elif __file__:
            absPath = os.path.dirname(os.path.abspath(__file__))
        useCustomTheme = True
        self.useCustomTheme = useCustomTheme
        self.absPath = absPath

        # SET HOME PAGE AND SELECT MENU
        # ///////////////////////////////////////////////////////////////
        widgets.stackedWidget.setCurrentWidget(widgets.home)
        widgets.btn_home.setStyleSheet(UIFunctions.selectMenu(widgets.btn_home.styleSheet()))

        # 初始化GenReport页面
        # 创建图纸展示画廊
        self.drawing_gallery = ImageGallery(self.ui.drawingShow)
        # 创建滚动区域
        drawing_scroll_area = QScrollArea()
        drawing_scroll_area.setWidgetResizable(True)
        drawing_scroll_area.setWidget(self.drawing_gallery)
        # 设置固定大小
        self.ui.drawingShow.setFixedSize(1500, 200)  # 设置固定大小
        self.ui.drawingShow.setLayout(QVBoxLayout())
        self.ui.drawingShow.layout().addWidget(drawing_scroll_area)

        # 创建实物图片展示画廊
        self.real_thing_gallery = ImageGallery(self.ui.picShow)
        # 创建滚动区域
        real_thing_scroll_area = QScrollArea()
        real_thing_scroll_area.setWidgetResizable(True)
        real_thing_scroll_area.setWidget(self.real_thing_gallery)
        # 设置固定大小
        self.ui.picShow.setFixedSize(1500, 200)  # 设置固定大小
        self.ui.picShow.setLayout(QVBoxLayout())
        self.ui.picShow.layout().addWidget(real_thing_scroll_area)

    # ----------------------------------------左侧菜单按钮------------------------------------------------------------------
    
    def buttonClick(self):
        # GET BUTTON CLICKED
        btn = self.sender()
        btnName = btn.objectName()


        if btnName == "btn_home":
            widgets.stackedWidget.setCurrentWidget(widgets.home)
            UIFunctions.resetStyle(self, btnName)
            btn.setStyleSheet(UIFunctions.selectMenu(btn.styleSheet()))
 
        if btnName == "btn_Flow":
            widgets.stackedWidget.setCurrentWidget(widgets.flow_check)
            UIFunctions.resetStyle(self, btnName)
            btn.setStyleSheet(UIFunctions.selectMenu(btn.styleSheet()))
            self.reflesh_Skulist()
            #更新所有完成按钮状态
            # self.update_all_done_buttons()

        if btnName == "btn_new":
            widgets.stackedWidget.setCurrentWidget(widgets.SkuList)
            UIFunctions.resetStyle(self, btnName)
            btn.setStyleSheet(UIFunctions.selectMenu(btn.styleSheet()))
            # 自动加载所有数据
            self.ui.reportTableView.load_data_from_db()
        
        if btnName == "btn_save":
            widgets.stackedWidget.setCurrentWidget(widgets.dict_setting) 
            UIFunctions.resetStyle(self, btnName)
            btn.setStyleSheet(UIFunctions.selectMenu(btn.styleSheet())) 
    
    # ----------------------------------------首页图表方法------------------------------------------------------------------
        """创建折线图"""    
    def create_line_chart(self):
        # 创建图表
        chart = QChart()
        chart.setTitle("数据趋势图")
        chart.setAnimationOptions(QChart.AnimationOption.SeriesAnimations)
        chart.setTheme(QChart.ChartTheme.ChartThemeLight)
        
        # 创建数据系列
        series = QLineSeries()
        series.setName("")
        series.setColor(QColor("#0256FF"))
        series.setPen(QPen(QColor("#0256FF"), 3))
        
        # 生成一些随机数据
        for i in range(10):
            series.append(i, random.randint(0, 100))
        
        # 添加数据系列到图表
        chart.addSeries(series)
        
        # 创建坐标轴
        axis_x = QValueAxis()
        axis_x.setTitleText("X轴")
        axis_x.setRange(0, 9)
        axis_x.setTickCount(10)
        axis_x.setLabelsColor(QColor("#333333"))
        axis_x.setTitleBrush(QBrush(QColor("#333333")))
        
        axis_y = QValueAxis()
        axis_y.setTitleText("Y轴")
        axis_y.setRange(0, 100)
        axis_y.setTickCount(11)
        axis_y.setLabelsColor(QColor("#333333"))
        axis_y.setTitleBrush(QBrush(QColor("#333333")))
        
        # 添加坐标轴到图表
        chart.addAxis(axis_x, Qt.AlignmentFlag.AlignBottom)
        chart.addAxis(axis_y, Qt.AlignmentFlag.AlignLeft)
        series.attachAxis(axis_x)
        series.attachAxis(axis_y)
        
        # 创建图表视图
        chart_view = QChartView(chart)
        chart_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 设置图表视图的样式
        chart_view.setStyleSheet("""
            QChartView {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 10px;
            }
            QChart {
                background-color: transparent;
            }
            QChart::title {
                color: #333333;
                font-size: 14px;
                font-weight: bold;
            }
        """)
        
        # 将图表视图添加到ColumFrame
        layout = QVBoxLayout(self.ui.ColumFrame)
        layout.addWidget(chart_view)

    """创建柱状图"""
    def create_bar_chart(self):
        # 创建图表
        chart = QChart()
        chart.setTitle("每月报告数量")
        chart.setAnimationOptions(QChart.AnimationOption.SeriesAnimations)
        chart.setTheme(QChart.ChartTheme.ChartThemeLight)
        
        # 创建数据系列
        series = QBarSeries()
        
        # 创建数据集合
        bar_set = QBarSet("报告数量")
        bar_set.setColor(QColor("#0256FF"))
        
        # 生成一些随机数据
        categories = ["一月", "二月", "三月", "四月", "五月", "六月"]
        for i in range(len(categories)):
            bar_set.append(random.randint(0, 100))
        
        # 添加数据集合到系列
        series.append(bar_set)
        
        # 添加数据系列到图表
        chart.addSeries(series)
        
        # 创建分类轴
        axis_x = QBarCategoryAxis()
        axis_x.append(categories)
        axis_x.setTitleText("月份")
        axis_x.setLabelsColor(QColor("#333333"))
        axis_x.setTitleBrush(QBrush(QColor("#333333")))
        
        # 创建数值轴
        axis_y = QValueAxis()
        axis_y.setTitleText("数值")
        axis_y.setRange(0, 100)
        axis_y.setTickCount(11)
        axis_y.setLabelsColor(QColor("#333333"))
        axis_y.setTitleBrush(QBrush(QColor("#333333")))
        
        # 添加坐标轴到图表
        chart.addAxis(axis_x, Qt.AlignmentFlag.AlignBottom)
        chart.addAxis(axis_y, Qt.AlignmentFlag.AlignLeft)
        series.attachAxis(axis_x)
        series.attachAxis(axis_y)
        
        # 创建图表视图
        chart_view = QChartView(chart)
        chart_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 设置图表视图的样式
        chart_view.setStyleSheet("""
            QChartView {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 10px;
            }
            QChart {
                background-color: transparent;
            }
            QChart::title {
                color: #333333;
                font-size: 14px;
                font-weight: bold;
            }
        """)
        
        # 将图表视图添加到LineChartFrame
        layout = QVBoxLayout(self.ui.LineChartFrame)
        layout.addWidget(chart_view)

    """创建饼图"""
    def create_pie_chart(self):
        # 创建图表
        chart = QChart()
        chart.setTitle("机种分布图")
        chart.setAnimationOptions(QChart.AnimationOption.SeriesAnimations)
        chart.setTheme(QChart.ChartTheme.ChartThemeLight)
        
        # 创建饼图系列
        series = QPieSeries()
        
        # 使用假数据替代数据库查询
        # country_data = self.db_manager.get_countryCount()
        country_data = [
            ("Butterworth ", 15),
            ("Calabasas ", 8),
            ("Kabini", 6),
            ("Queen City", 4),
            ("Berry Creek", 3)
        ]
        
        # 定义颜色列表
        colors = [QColor("#FF6384"), QColor("#36A2EB"), QColor("#FFCE56"), 
                 QColor("#4BC0C0"), QColor("#9966FF"), QColor("#FF9F40"),
                 QColor("#8AC24A"), QColor("#9C27B0"), QColor("#E91E63")]
        
        # 添加数据切片
        for i, (country, count) in enumerate(country_data):
            if count > 0:  # 只显示有数据的国家
                slice = series.append(f"{country} ({count})", count)
                slice.setColor(colors[i % len(colors)])  # 循环使用颜色列表
                slice.setLabelVisible(True)
                slice.setLabelPosition(QPieSlice.LabelPosition.LabelOutside)
                slice.setExploded(False)
                slice.setLabelColor(QColor("#333333"))
                slice.setLabelFont(QFont("Microsoft YaHei", 9))
        
        # 添加系列到图表
        chart.addSeries(series)
        
        # 创建图表视图
        chart_view = QChartView(chart)
        chart_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 设置图表视图的样式
        chart_view.setStyleSheet("""
            QChartView {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 10px;
            }
            QChart {
                background-color: transparent;
            }
            QChart::title {
                color: #333333;
                font-size: 14px;
                font-weight: bold;
            }
        """)
        
        # 将图表视图添加到PiesChartsFrame
        layout = QVBoxLayout(self.ui.PiesChartsFrame)
        layout.addWidget(chart_view)

    # """创建日历视图"""
    # def create_calendar(self):
    #     # 创建日历控件
    #     calendar = QCalendarWidget()
    #     calendar.setGridVisible(True)
        
    #     # 设置日历样式
    #     calendar.setStyleSheet("""
    #         QCalendarWidget {
    #             background-color: #ffffff;
    #             border: 1px solid #e0e0e0;
    #             border-radius: 8px;
    #             padding: 5px;
    #         }
    #         QCalendarWidget QToolButton {
    #             height: 30px;
    #             width: 100px;
    #             color: #0256FF;
    #             font-size: 14px;
    #             font-weight: bold;
    #             background-color: #ffffff;
    #             border: none;
    #         }
    #         QCalendarWidget QToolButton:hover {
    #             background-color: #f0f0f0;
    #         }
    #         QCalendarWidget QMenu {
    #             width: 150px;
    #             left: 20px;
    #             background-color: #ffffff;
    #             border: 1px solid #e0e0e0;
    #             border-radius: 4px;
    #         }
    #         QCalendarWidget QSpinBox {
    #             width: 50px;
    #             font-size: 12px;
    #             color: #0256FF;
    #             background-color: #ffffff;
    #             selection-background-color: #0256FF;
    #             selection-color: white;
    #             border: 1px solid #e0e0e0;
    #             border-radius: 4px;
    #         }
    #         QCalendarWidget QAbstractItemView:enabled {
    #             font-size: 12px;
    #             color: #333333;
    #             background-color: #ffffff;
    #             selection-background-color: #0256FF;
    #             selection-color: white;
    #             border: none;
    #         }
    #         QCalendarWidget QAbstractItemView:disabled {
    #             color: #cccccc;
    #         }
    #         QCalendarWidget QWidget#qt_calendar_navigationbar {
    #             background-color: #ffffff;
    #             border-bottom: 1px solid #e0e0e0;
    #         }
    #         QCalendarWidget QWidget#qt_calendar_navigationbar QToolButton {
    #             background-color: #ffffff;
    #             border: none;
    #             color: #0256FF;
    #         }
    #         QCalendarWidget QWidget#qt_calendar_navigationbar QToolButton:hover {
    #             background-color: #f0f0f0;
    #         }
    #         /* 今日日期样式 */
    #         QCalendarWidget QAbstractItemView:enabled:selected {
    #             background-color: #0256FF;
    #             color: white;
    #             border-radius: 4px;
    #         }
    #         /* 今日日期高亮 */
    #         QCalendarWidget QAbstractItemView:enabled:!selected:has-focus {
    #             background-color: #e6f0ff;
    #             color: #0256FF;
    #             border-radius: 4px;
    #         }
    #         /* 周末日期样式 */
    #         QCalendarWidget QAbstractItemView:enabled:!selected:!has-focus {
    #             color: #666666;
    #         }
    #         /* 周末日期选中样式 */
    #         QCalendarWidget QAbstractItemView:enabled:selected:has-focus {
    #             background-color: #0256FF;
    #             color: white;
    #             border-radius: 4px;
    #         }
    #         /* 周末日期高亮样式 */
    #         QCalendarWidget QAbstractItemView:enabled:!selected:has-focus {
    #             background-color: #e6f0ff;
    #             color: #0256FF;
    #             border-radius: 4px;
    #         }
    #     """)
        
    #     # 设置当前日期
    #     calendar.setSelectedDate(QDate.currentDate())
        
    #     # 将日历控件添加到CalendarFrame
    #     layout = QVBoxLayout(self.ui.CalendarFrame)
    #     layout.addWidget(calendar)

    # -----------------------------------------流程按钮方法------------------------------------------------------------------
    
    def check_status(self, item,des):
        # 检查是否已选择SKU
        if not self.selectedSku:
            CustomMessageBox.warning(None, "提示", "请先选择一个SKU！")
            return False
        # 检查是否上个流程已完成
        flow_status = self.db_manager.select_skuDic_list()
        status_completed = False
        
        for status_dict in flow_status:
            if self.selectedSku in status_dict:
                sku_status = status_dict[self.selectedSku]
                if sku_status.get(item) == '1':
                    status_completed = True
                break
        
        if not status_completed:
            CustomMessageBox.warning(None, "警告", des)
            return False
        return True
    
    def btnClick_BomCheck(self):
        # 检查是否已选择SKU
        if not self.selectedSku:
            CustomMessageBox.warning(None, "提示", "请先选择一个SKU！")
            return
        CustomMessageBox.info(None, "提示", "当前流程需要在Power Automate 上完成！")
        self.ui.BomShowLabel.setText(f"BOM比对执行的SKU是：{self.selectedSku}")
        self.init_bom_file_list()
        widgets.stackedWidget.setCurrentWidget(widgets.BomCheck)

    def btnClick_BomReturn(self):
        #清空Producttext和Languagetext的内容
        self.ui.Producttext.clear()
        self.ui.Languagetext.clear()
        widgets.stackedWidget.setCurrentWidget(widgets.flow_check)
    
    def btnClick_PicReturn(self):
        widgets.stackedWidget.setCurrentWidget(widgets.flow_check)
    
    def btnClick_PicDownload(self):
        # 检查是否已选择SKU
        if not self.selectedSku:
            CustomMessageBox.warning(None, "提示", "请先选择一个SKU！")
            return
        if not self.check_status('bom_check','请先在Power Automate 上完成BOM比对流程！'):
            return
        CustomMessageBox.info(None, "提示", "当前流程需要在Power Automate 上完成！")
        self.ui.pic_info.setText(f"图纸下载执行的SKU是：{self.selectedSku}")
        widgets.stackedWidget.setCurrentWidget(widgets.PicDownload)
        # 加载列表

        self.ui.tableview.load_data_from_db(self.selectedSku)
    def btnClick_PicCheck(self):

        try:               
            if not self.check_status('pic_download','请先在Power Automate 上完成图纸下载流程！'):
                return
            
            from modules.OtherView.TakePicView import TakePicWindow
            
            # 确保camera_window被正确初始化
            if not hasattr(self, 'camera_window'):
                self.camera_window = None
                
            # 如果窗口已存在但被关闭了，重新创建
            if self.camera_window is not None and not self.camera_window.isVisible():
                self.camera_window = None
            
            # 创建新窗口
            if self.camera_window is None:
                try:
                    self.camera_window = TakePicWindow(self, self.selectedSku)
                    if self.camera_window is None:
                        raise Exception("创建相机窗口失败")
                        
                    # 设置窗口标题
                    self.camera_window.setWindowTitle(f"相机拍照 - {self.selectedSku}")
                    # 设置模态窗口
                    self.camera_window.setWindowModality(Qt.WindowModality.NonModal)
                except Exception as e:
                    CustomMessageBox.warning(None, "错误", f"创建相机窗口失败：{str(e)}")
                    self.camera_window = None
                    return
                
            # 显示窗口
            try:
                self.camera_window.show()
                self.camera_window.raise_()
                self.camera_window.activateWindow()
                self.ui.btnFlow.pic_check.setChecked(True)
            except Exception as e:
                CustomMessageBox.warning(None, "错误", f"显示相机窗口失败：{str(e)}")
                self.camera_window = None
            
        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"打开相机窗口失败：{str(e)}")
            self.camera_window = None
    
    def btnClick_GenReport(self):
        if not self.check_status('pic_check','请先完成拍照比对流程！'):
            return
        self.init_gen_report()
        widgets.stackedWidget.setCurrentWidget(widgets.GenReport)
        self.real_thing_gallery.load_real_thing_images(self.selectedSku)
        # 读取references_label的值
        self.load_references_label()
    
    def drawEnd(self):
        # 检查模板是否上传，是否截图
        if not self.db_manager.check_template(self.selectedSku):
            CustomMessageBox.warning(None, "警告", "物料信息还未上传!")
            return
        if not os.path.exists(os.path.join(ProjectSettings.PIC_DOWNLOAD_PDF_PATH, self.selectedSku)):
            CustomMessageBox.warning(None, "警告", "还未上传任何截图!")
            return
        """确认图纸检查流程完成"""
        reply = CustomMessageBox.question(
            None,
            "确认",
            "该流程确认完成？",
            CustomMessageBox.StandardButton.Yes | CustomMessageBox.StandardButton.No,
            CustomMessageBox.StandardButton.No
        )
        
        if reply == CustomMessageBox.StandardButton.Yes:
            # 更新流程状态
            self.db_manager.update_flow_status(self.selectedSku, 'pic_download', '1')
            # 更新按钮样式
            self.ui.btnFlow.add_BtnDoneStyle(self.ui.btnFlow.pic_download)
            # 返回主流程页面
            widgets.stackedWidget.setCurrentWidget(widgets.flow_check)
            
    def bomEnd(self):
        """确认BOM检查流程完成"""
        reply = CustomMessageBox.question(
            None,
            "确认",
            "该流程确认完成？",
            CustomMessageBox.StandardButton.Yes | CustomMessageBox.StandardButton.No,
            CustomMessageBox.StandardButton.No
        )
        
        if reply == CustomMessageBox.StandardButton.Yes:
           
            if not self.read_Product_Language_BomPath():               
                return
            # 更新流程状态
            self.db_manager.update_flow_status(self.selectedSku, 'bom_check', '1')

            # 更新按钮样式
            self.ui.btnFlow.add_BtnDoneStyle(self.ui.btnFlow.bom_check)
            # 返回主流程页面
            widgets.stackedWidget.setCurrentWidget(widgets.flow_check)
            
        
    def read_Product_Language_BomPath(self):
        """从前端文本框读取产品语言和BOM路径信息
        
        Returns:
            bool: 读取成功返回True，否则返回False
        """
        try:
            # 检查是否已选择SKU
            if not self.selectedSku:
                CustomMessageBox.warning(None, "警告", "请先选择一个SKU！")
                return False
            
            # 从前端文本框获取产品名和语言信息
            product = self.ui.Producttext.toPlainText().strip()
            language = self.ui.Languagetext.toPlainText().strip()
            
            print(f"DEBUG: 从文本框获取的产品名: '{product}'")  # 调试输出
            print(f"DEBUG: 从文本框获取的语言: '{language}'")  # 调试输出
            
            # 检查产品名是否为空
            if not product:
                CustomMessageBox.warning(None, "警告", "请输入产品名！")
                return False
            
            # 检查语言是否为空
            if not language:
                CustomMessageBox.warning(None, "警告", "请输入语言！")
                return False

            # 获取BOM文件路径
            sku_dir = os.path.join(ProjectSettings.BOM_CHECK_PATH, self.selectedSku)
            print(f"DEBUG: BOM文件路径: {sku_dir}")
            if not os.path.exists(sku_dir):
                CustomMessageBox.warning(None, "警告", f"当前还未使用RPA下载BOM")
                return False
                
            # 获取各种BOM文件
            print(f"DEBUG: 当前SKU: {self.selectedSku}")
            ckm1_bom_path = [file for file in os.listdir(sku_dir) if file.startswith(f"{self.selectedSku} CKM1 BOM")]
            msft_bom_path = [file for file in os.listdir(sku_dir) if file.startswith(f"{self.selectedSku} MSFT BOM")]
            fai_observations_path = [file for file in os.listdir(sku_dir) if file.startswith(f"{self.selectedSku} MSFT BOM & CKM1 BOM Comparison result") and file.endswith(".xlsx")]
            print(f"DEBUG: 找到的CKM1 BOM文件: {ckm1_bom_path}")
            print(f"DEBUG: 找到的MSFT BOM文件: {msft_bom_path}")
            print(f"DEBUG: 找到的FAI观察文件: {fai_observations_path}")
            # 检查是否找到必要的文件
            if not ckm1_bom_path or not msft_bom_path:
                CustomMessageBox.warning(None, "警告", "未找到必要的BOM文件")
                return False

            # 更新数据库
            if self.db_manager.Insert_BOM_Path(self.selectedSku, product, language, ckm1_bom_path, msft_bom_path, fai_observations_path):
                CustomMessageBox.info(None, "成功", "BOM路径信息已更新")
                return True
            else:
                CustomMessageBox.warning(None, "错误", "更新BOM路径信息失败")
                return False
                
        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"读取BOM路径信息时出错: {str(e)}")
            print(f"错误: {str(e)}")
            return False

    def update_all_done_buttons(self):
        # 更新所有完成按钮状态

        flow_status = self.db_manager.select_skuDic_list()
        for status_dict in flow_status:
            if self.selectedSku in status_dict:
                sku_status = status_dict[self.selectedSku]
                for key, value in sku_status.items():
                    if key == 'bom_check' and value == '1':
                        # 按钮置灰,不执行点击事件
                        self.ui.bomEndBtn.setDisabled(True)
                        self.ui.bomEndBtn.clicked.disconnect()
                    if key == 'pic_download' and value == '1':
                        self.ui.drawEndBtn.setDisabled(True)
                        self.ui.drawEndBtn.clicked.disconnect()
        
    # -----------------------------------------添加SKU方法------------------------------------------------------------------
    
    def addMoreSku(self):
        self.ui_addSku.show()
        self.reflesh_Skulist()

    def reflesh_Skulist(self):
        SKU_list = self.db_manager.select_sku_list()
        self.ui.skuList.clear()
        self.ui.skuList.add_Items_sku(SKU_list)
        if self.ui.skuList.count() > 0:
            self.sku_clicked(self.ui.skuList.item(0) )

    def filter_list(self):
        # 清空 QListWidget
        self.ui.skuList.clear()
        # 重新匹配
        text = self.ui.skuSearchLine.text()
        if text == "":
            self.reflesh_Skulist()
            return
        SKU_list = self.db_manager.select_sku_list()
        for item in SKU_list:
            if text.lower() in item.lower():  # 模糊匹配
                if not self.ui.skuList.findItems(item, Qt.MatchExactly):
                    self.ui.skuList.add_Item_sku(item)

    def sku_clicked(self, item):
        self.ui.btnFlow.Init_BtnStyle()
        self.selectedSku=item.text()
        key = item.text()
        flow_status = self.db_manager.select_skuDic_list()
        for statusDict in flow_status:
            value = statusDict.get(key)
            if value:
                # 根据不同的状态值设置对应按钮的样式
                if value.get('bom_check') == '1':
                    self.ui.btnFlow.add_BtnDoneStyle(self.ui.btnFlow.bom_check)
                if value.get('pic_download') == '1':
                    self.ui.btnFlow.add_BtnDoneStyle(self.ui.btnFlow.pic_download)
                if value.get('pic_check') == '1':
                    self.ui.btnFlow.add_BtnDoneStyle(self.ui.btnFlow.pic_check)
                if value.get('gen_report') == '1':
                    self.ui.btnFlow.add_BtnDoneStyle(self.ui.btnFlow.gen_report)
        
        # 加载图纸到图纸画廊
        self.drawing_gallery.load_images(self.selectedSku)
        # 加载实物图片到实物图片画廊
        self.real_thing_gallery.load_real_thing_images(self.selectedSku)
        # 加载BOM文件列表
    
    def delete_sku_from_database(self, sku_text):
        """从数据库中删除SKU"""
        try:
            # 从数据库中删除SKU
            self.db_manager.delete_sku(sku_text)
            print(f"已从数据库中删除SKU: {sku_text}")
            
            # 如果删除的是当前选中的SKU，清空选择
            if hasattr(self, 'selectedSku') and self.selectedSku == sku_text:
                self.selectedSku = None
                self.ui.btnFlow.Init_BtnStyle()
                # 清空图片画廊
                self.drawing_gallery.clear_images()
                self.real_thing_gallery.clear_images()
            
            # 显示删除成功消息
            CustomMessageBox.info(None, "删除成功", f"已成功删除SKU: {sku_text}")
            
        except Exception as e:
            print(f"删除SKU时出错: {e}")
            CustomMessageBox.warning(None, "删除失败", f"删除SKU时出错: {str(e)}")
    
    # -----------------------------------------BOM比对页面方法------------------------------------------------------------------
    
    def init_bom_file_list(self):
        """初始化BOM文件列表"""
        # 设置列表视图的样式
        self.ui.BOMFileList.setStyleSheet("""
            QListView {
                background-color: #f5f5f5;
                border: 1px solid #e0e0e0;
                border-radius: 4px;
                padding: 5px;
                font-size: 12px;
            }
            QListView::item {
                height: 30px;
                padding: 5px;
                margin: 2px;
                border-radius: 4px;
                color: #333333;
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
            }
            QListView::item:hover {
                background-color: #e6f0ff;
                border: 1px solid #0256FF;
                color: #0256FF;
            }
            QListView::item:selected {
                background-color: #0256FF;
                color: white;
                border: 1px solid #0256FF;
            }
            QListView::item:selected:active {
                background-color: #0256FF;
                color: white;
            }
            QListView::item:selected:!active {
                background-color: #0256FF;
                color: white;
            }
        """)
        
        # 创建模型
        self.bom_file_model = QStandardItemModel()
        self.ui.BOMFileList.setModel(self.bom_file_model)
        
        # 连接点击事件
        self.ui.BOMFileList.clicked.connect(self.on_bom_file_clicked)
        
        # 加载Excel文件
        self.load_bom_files()
        
    def load_bom_files(self):
        """加载BOM文件夹中的Excel文件"""
        try:
            # 检查是否选择了SKU
            if not self.selectedSku:
                return
                
            # 检查模型是否已初始化
            if not hasattr(self, 'bom_file_model'):
                self.init_bom_file_list()
                return
                
            # 构建文件夹路径
            folder_path = os.path.join(ProjectSettings.BOM_CHECK_PATH, self.selectedSku)
            
            # 检查文件夹是否存在
            if not os.path.exists(folder_path):
                return
                
            # 获取所有Excel文件
            excel_files = []
            for file in os.listdir(folder_path):
                if file.endswith(('.xlsx', '.xls')):
                    excel_files.append(file)
                    
            # 清空列表
            self.bom_file_model.clear()
            
            # 添加文件到列表
            for file in excel_files:
                item = QStandardItem(file)
                self.bom_file_model.appendRow(item)
                
        except Exception as e:
            print(f"加载BOM文件列表时出错: {str(e)}")
            
    def on_bom_file_clicked(self, index):
        """处理BOM文件点击事件"""
        try:
            # 获取选中的文件名
            file_name = self.bom_file_model.data(index)
            
            # 构建完整的文件路径
            file_path = os.path.join(ProjectSettings.BOM_CHECK_PATH, self.selectedSku, file_name)
            
            # 检查文件是否存在
            if os.path.exists(file_path):
                # 使用系统默认程序打开文件
                os.startfile(file_path)
            else:
                CustomMessageBox.warning(None, "警告", f"文件不存在: {file_path}")
                
        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"打开文件时出错: {str(e)}")

    # -----------------------------------------图纸上传页面方法------------------------------------------------------------------
    
    def uploadCheckFile(self):
        try:
            # 检查是否已选择SKU
            if not self.selectedSku:
                CustomMessageBox.warning(None, "警告", "请先选择一个SKU！")
                return
                
            # 检查SKU是否已存在于drawing表中
            if self.db_manager.check_sku_exists_in_drawing(self.selectedSku):
                reply = CustomMessageBox.question(
                    None,
                    "提示",
                    f"SKU {self.selectedSku} 已经上传过，是否继续上传，重新上传原有数据和图片都将全部删除？",
                    CustomMessageBox.StandardButton.Yes | CustomMessageBox.StandardButton.No,
                    CustomMessageBox.StandardButton.No
                )
                if reply == CustomMessageBox.StandardButton.No:
                    return
                if reply == CustomMessageBox.StandardButton.Yes:
                    # 数据库清除数据
                    self.db_manager.delete_data_by_sku(self.selectedSku)
                    self.db_manager.delete_pic_by_sku(self.selectedSku)
                    # 删除物料文件
                    BOM_BASE_PATH = os.path.join(ProjectSettings.BOM_CHECK_PATH, self.selectedSku)
                    BOM_DIR = os.path.join(BOM_BASE_PATH,self.selectedSku+"物料.xlsx")
                    os.remove(BOM_DIR)
                    # 删除文件夹
                    SCREENSHOT_DIR = os.path.join(ProjectSettings.REALPIC_PATH, self.selectedSku, "SCREENSHOT")
                    self.delete_folder(SCREENSHOT_DIR)
                     # 加载列表
                    self.ui.tableview.load_data_from_db(self.selectedSku)
                    # 加载图片画廊
                    self.image_gallery.load_images(self.selectedSku)

            # 打开文件选择对话框
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "选择Excel文件",
                "",
                "Excel文件 (*.xlsx *.xls)"
            )

            if not file_path:  # 用户取消选择
                return


            # 导入Excel数据到数据库
            if self.db_manager.import_excel_to_drawing(self.db_path, self.selectedSku, file_path):
                
                CustomMessageBox.info(None, "成功", "数据已成功导入数据库")
                #将文件file_path复制到BOM_BASE_PATH
                BOM_BASE_PATH = os.path.join(ProjectSettings.BOM_CHECK_PATH, self.selectedSku)
                BOM_DIR = os.path.join(BOM_BASE_PATH,self.selectedSku+"物料.xlsx")
                shutil.copy(file_path, BOM_DIR)
                # 刷新表格显示
                self.ui.tableview.load_data_from_db(self.selectedSku)
            else:
                CustomMessageBox.warning(None, "错误", "导入数据失败")

        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"导入数据时出错: {str(e)}")
            print(f"导入数据时出错: {str(e)}")

    def delete_folder(self, folder_path):
        """
        功能描述：删除指定文件夹及其内容
        参数说明：folder_path - 要删除的文件夹路径
        返回值说明：无
        异常说明：无
        """
        try:
            shutil.rmtree(folder_path)
        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"删除文件夹时出错: {str(e)}")

    def downloadMB(self):
        WL_PATH = ProjectSettings.WULIAO_PATH
        # 下载这个文件到当前电脑桌面
        desktop_path = os.path.expanduser("~/Desktop")
        file_name = os.path.basename(WL_PATH)
        shutil.copy(WL_PATH, os.path.join(desktop_path, file_name))
        CustomMessageBox.info(None, "成功", "物料模板已下载到桌面")
        
    
    # -----------------------------------------报告生成页面方法------------------------------------------------------------------
    
    def init_gen_report(self):
        """初始化生成报告页面的UI和数据"""
        models = self.db_manager.load_model_data()
        self.ui.model_comboBox.addItems(models)
        
        # # 加载产品名数据
        
        self.ui.buildFlow.setText(f"当前生成的是{self.selectedSku}的报告")
        
        # 从report_generate表里根据sku拿到product_name,country
        product,country=self.db_manager.get_product_name_country(self.selectedSku)
        print(product,country)
        # 设置product_name和country的值
        
        self.ui.country_label.setText(country)
        self.ui.product_label.setText(product)
        
        # 创建references_textEdit替换references_label
        self.references_textEdit = QTextEdit(self.ui.referencesInfo)
        self.references_textEdit.setObjectName("references_textEdit")
        self.references_textEdit.setReadOnly(True)  # 设置为只读
        self.references_textEdit.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)  # 设置垂直滚动条
        self.references_textEdit.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)  # 设置水平滚动条
        self.references_textEdit.setStyleSheet("color: rgb(0, 85, 255);")  # 设置字体颜色为蓝色
        
        # 获取referencesInfo的布局referencesInfo
        if not self.ui.referencesInfo.layout():
            layout = QVBoxLayout(self.ui.referencesInfo)
        else:
            layout = self.ui.referencesInfo.layout()
            
        # 清除布局中的所有控件
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
                
        # 添加references_textEdit到布局
        layout.addWidget(self.references_textEdit)
        
        # 读取references_label的值
        self.load_references_label()

    def on_model_changed(self, model_name):
        """当机种改变时更新标签"""
        self.ui.model_label.setText(model_name)

    def generate_report(self):
        # 检查是否选择了SKU
        if not self.selectedSku:
            CustomMessageBox.warning(None, "警告", "请先选择一个SKU！")
            return

        # 获取标签值
        model_name = self.ui.model_label.text()
        product_name = self.ui.product_label.text()
        sn = self.ui.SN_label.text()
        country = self.ui.country_label.text()
        faiDate = datetime.now().strftime("%Y.%m.%d")
        reference_info = self.references_textEdit.toPlainText()
        
        # 检查必要的值是否为空
        if not model_name:
            CustomMessageBox.warning(None, "警告", "请先选择机种！")
            return
        # 根据model_name获取model_name_path
        template_path = self.db_manager.get_template_path(model_name)

        # 检查模板文件是否存在
        if not os.path.exists(template_path):
            CustomMessageBox.warning(None, "警告", "模板文件不存在！")
            return
        
        from docx import Document
        from docx.shared import Inches
        
       
        # 假设self.selectedSku、self.report_path、self.output_path已在类中定义或可获取
        output_path = ProjectSettings.REPORT_PATH + self.selectedSku + ".docx"  

        try:
            generate_report_with_images(self.selectedSku, template_path, output_path, self.db_manager)
             # 打开模板文档
            doc = Document(output_path)
            
            # 替换文档中的占位符
            for paragraph in doc.paragraphs:
                if "{{sku}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{sku}}", self.selectedSku)
                if "{{product_name}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{product_name}}", product_name)
                if "{{serial_number}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{serial_number}}", sn)
                if "{{fai_date}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{fai_date}}", faiDate)
                if "{{country}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{country}}", country)
                if "{{reference_info}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{reference_info}}", reference_info)
                
            # 处理表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "{{sku}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{sku}}", self.selectedSku)
                            if "{{product_name}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{product_name}}", product_name)
                            if "{{serial_number}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{serial_number}}", sn)
                            if "{{fai_date}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{fai_date}}", faiDate)
                            if "{{country}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{country}}", country)
                            if "{{reference_info}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{reference_info}}", reference_info)
            # 保存文档
            doc.save(output_path)  
            # 更改数据库状态
            self.db_manager.update_flow_status(self.selectedSku, 'gen_report', '1')
                # 更新所有完成按钮状态
            self.db_manager.update_flow_status(self.selectedSku, 'Status', '1')
                # 更新报告状态
            self.db_manager.insert_report_path(self.selectedSku, reference_info,sn,faiDate,output_path, self.selectedSku+".docx")
            self.sku_clicked(self.ui.skuList.item(0) )
            CustomMessageBox.information(None, "成功", f"报告已生成：{output_path}")
            os.startfile(output_path)
                
        except Exception as e:
            CustomMessageBox.error(None, "错误", f"生成报告失败: {str(e)}")
        

    def update_sn_label(self):
        """更新SN标签内容"""
        sn_text = self.ui.SN_textEdit.toPlainText()
        self.ui.SN_label.setText(sn_text)

    def load_references_label(self):
        """从数据库中读取references信息"""
        try:
            # 检查是否选择了SKU
            if not self.selectedSku:
                return
                
            # 获取数据库的值
            references = self.db_manager.get_references_by_sku(self.selectedSku)
            
            # 将references_textEdit的值设置为获取到的reference_info
            if references:
                self.references_textEdit.setText("\n".join(references))
            else:
                self.references_textEdit.clear()
            
        except Exception as e:
            print(f"加载references信息失败: {e}")
           
    def btnClick_genBack(self):
        widgets.stackedWidget.setCurrentWidget(widgets.flow_check)
        self.buttonClick

    # -----------------------------------------做好的SKU页面------------------------------------------------------------------
    def redoSkuTable(self):
        """重置表格显示所有数据"""
        try:
            # 获取表格实例
            table = self.ui.reportTableView
            # 重新加载所有数据
            table.load_data_from_db()
        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"重置表格时出错: {str(e)}")

    def skuTableQuery(self):
        """根据SKU和日期范围查询数据"""
        try:
            # 获取表格实例
            table = self.ui.reportTableView
            
            # 获取SKU输入
            sku = self.ui.skuInput.toPlainText().strip()
            
            # 获取日期范围
            start_date = self.ui.startDateEdit.date().toString("yyyy-MM-dd")
            end_date = self.ui.endDateEdit.date().toString("yyyy-MM-dd")
            
            # 如果同时输入了SKU和日期范围，使用组合查询
            if sku and start_date and end_date:
                table.query_by_sku_and_date(sku, start_date, end_date)
            # 如果只输入了SKU，按SKU查询
            elif sku:
                table.query_by_sku(sku)
            # 如果只选择了日期范围，按日期查询
            elif start_date and end_date:
                table.query_by_date(start_date, end_date)
            # 如果都没有输入，显示所有数据
            else:
                table.load_data_from_db()
                
        except Exception as e:
            CustomMessageBox.warning(None, "错误", f"查询数据时出错: {str(e)}")

    # -----------------------------------------主UI事件方法------------------------------------------------------------------
    def center_window(self):
        """将窗口居中显示"""
        # 获取屏幕的中心点
        screen_geometry = QApplication.primaryScreen().availableGeometry()
        screen_center = screen_geometry.center()
        # 计算窗口的位置
        window_geometry = self.frameGeometry()
        window_geometry.moveCenter(screen_center)
        self.move(window_geometry.topLeft())

    def mousePressEvent(self, event):
        # SET DRAG POS WINDOW
        self.dragPos = event.globalPos()

        # PRINT MOUSE EVENTS
        if event.buttons() == Qt.LeftButton:
            print('Mouse click: LEFT CLICK')
        if event.buttons() == Qt.RightButton:
            print('Mouse click: RIGHT CLICK')
    # RESIZE EVENTS
    # ///////////////////////////////////////////////////////////////
    def resizeEvent(self, event):
        # Update Size Grips
        UIFunctions.resize_grips(self)

    def closeEvent(self, event):
        # 确保主窗口获得焦点
        self.raise_()
        self.activateWindow()
        
        # 弹出消息框，询问用户是否确定关闭窗口
        reply = CustomMessageBox.question(
            None,
            '关闭窗口',
            '确定要关闭窗口吗？',
            CustomMessageBox.StandardButton.Yes | CustomMessageBox.StandardButton.No,
            CustomMessageBox.StandardButton.No
        )

        # 如果用户选择确定
        if reply == CustomMessageBox.StandardButton.Yes:
            # # 取消计时器
            # if self.timer is not None:
            #     self.timer.cancel()
            # 关闭相机窗口
            if hasattr(self, 'camera_window') and self.camera_window is not None:
                self.camera_window.close()
                self.camera_window = None
            # 退出应用程序
            QApplication.instance().quit()
        else:
            # 忽略关闭事件，保持窗口打开
            event.ignore()

           

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("icon.ico"))
    window = MainWindow()
    sys.exit(app.exec())
